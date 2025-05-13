from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE


def create_project_report(filepath):
    # Create a new Word document
    doc = Document()

    # Title
    title = doc.add_heading("Student Dashboard - Project Report", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Overview
    doc.add_heading("Overview", level=2)
    doc.add_paragraph(
        "The Student Dashboard is a web application built using the Streamlit framework. "
        "It provides an interactive platform for managing and visualizing student data, "
        "including names, grades, and performance levels. The application is designed with "
        "a modern dark theme and includes various features for data input, visualization, "
        "and statistical analysis."
    )

    # Features
    doc.add_heading("Features", level=2)
    features = [
        "Dark Theme and Styling: A custom dark theme is applied using CSS for a visually appealing interface.",
        "Header and Footer: A header image is included to make the app more engaging, and a footer with branding is added.",
        "Data Input: Users can input student details (name, grade, and performance) through text fields.",
        "Data Display: Displays all student records in a sorted table (sorted by performance and grade).",
        "Statistics: Calculates and displays key statistics for grades (Mean, Median, Mode).",
        "Visualizations: Includes a histogram for grades and a pie chart for performance distribution.",
        "Individual Sorting: Allows sorting data by name, grade, or performance with interactive visuals.",
    ]
    for feature in features:
        doc.add_paragraph(f"- {feature}")

    # Technical Highlights
    doc.add_heading("Technical Highlights", level=2)
    highlights = [
        "Streamlit Framework: Used for creating interactive web applications.",
        "Database Integration: Interacts with a database through the Student class.",
        "Plotly for Visualizations: Creates interactive and visually appealing charts.",
        "CSS for Custom Styling: Applies a dark theme and improves the app's appearance.",
    ]
    for highlight in highlights:
        doc.add_paragraph(f"- {highlight}")

    # Define a custom style for code snippets
    styles = doc.styles
    code_style = styles.add_style("CodeStyle", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(10)

    # Insert a code snippet for the dark theme CSS
    doc.add_heading("Dark Theme CSS Code", level=2)
    css_code = """
    <style>
    body {
        background-color: #121212;
        color: #ffffff;
    }
    .stButton>button {
        background-color: #1f77b4;
        color: white;
        border-radius: 8px;
        padding: 10px 20px;
    }
    .stButton>button:hover {
        background-color: #ff7f0e;
        color: black;
    }
    .stDataFrame {
        background-color: #1e1e1e;
        color: white;
    }
    </style>
    """
    doc.add_paragraph(css_code, style="CodeStyle")  # Use the custom "CodeStyle"

    # Strengths
    doc.add_heading("Strengths", level=2)
    strengths = [
        "User-Friendly Interface: Intuitive and easy to use.",
        "Interactive Visualizations: Enhances data exploration and understanding.",
        "Modern Design: Dark theme and custom styling make the app visually appealing.",
        "Comprehensive Features: Combines data management, statistics, and visualization.",
    ]
    for strength in strengths:
        doc.add_paragraph(f"- {strength}")

    # Suggestions for Further Improvements
    doc.add_heading("Suggestions for Further Improvements", level=2)
    improvements = [
        "Error Handling: Add error messages for invalid inputs and handle database connection errors.",
        "Performance Optimization: Cache database queries to improve performance for large datasets.",
        "Export Options: Allow users to export the displayed data (e.g., as a CSV file).",
        "Authentication: Add user authentication to restrict access to sensitive data.",
        "Mobile Responsiveness: Optimize the layout for better usability on mobile devices.",
    ]
    for improvement in improvements:
        doc.add_paragraph(f"- {improvement}")

    # Conclusion
    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(
        "The Student Dashboard is a well-designed and feature-rich application for managing "
        "and visualizing student data. It effectively combines functionality, aesthetics, and "
        "interactivity, making it a valuable tool for educators and administrators. With minor "
        "enhancements, it can become even more robust and user-friendly."
    )

    # Save the document
    doc.save(filepath)
    print(f"Report saved to {filepath}")


# Generate the report
create_project_report(
    "c:\\Users\\saher\\Desktop\\workshop\\algorithm_project\\Student_Dashboard_Report.docx"
)
